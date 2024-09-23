import sys
import time
# initial prompt
import tkinter as tk
from pathlib import Path
from tkinter import simpledialog
from typing import List, Tuple

import PyPDF2
from docx import Document
# doc import
from docx.opc.exceptions import PackageNotFoundError


# prep functions - setup git
def check_git_pull() -> int:
    """
    Prompts the user to confirm if they have pulled the latest changes from GitHub.
    If the user confirms, asks for the lesson number and returns it.
    If the user has not pulled the latest changes, exits the script after displaying a reminder.

    Returns:
        int: The lesson number input by the user.
    """
    root = tk.Tk()
    root.withdraw()  # Hide the root window

    # Prompt the user to confirm if they have pulled the latest changes
    response = simpledialog.askstring("Input", "Have you pulled the latest changes to your lesson slides (y/n)?")

    if response is None or response.lower() not in ['y', 'yes']:
        print("""Please update your most recent slides before proceeding.\ne.g. git pull from your Overleaf/LaTeX repo""")
        root.destroy()
        return None  # Instead of exiting, return None to indicate termination

    # If the user confirms, ask for the lesson number
    # Loop to ensure a valid lesson number is entered
    while True:
        lesson_no = simpledialog.askstring("Input", "What lesson number are you teaching?")
        if lesson_no is None:
            print("Lesson number entry was cancelled. Exiting.")
            root.destroy()
            return None  # Exit if the user cancels the dialog
        if lesson_no.isdigit():
            root.destroy()
            return int(lesson_no)  # Return the valid lesson number
        else:
            simpledialog.messagebox.showerror("Error", "Please enter a valid lesson number (numeric).")


# load docs funcs
def find_lesson_indices(syllabus: List[str], current_lesson: int) -> Tuple[int, int, int]:
    """
    Finds the indices of the lessons in the syllabus content.

    Args:
        syllabus (List[str]): A list of strings where each string represents a line in the syllabus document.
        current_lesson (int): The lesson number for which to find surrounding lessons.

    Returns:
        Tuple[int, int, int]: The indices of the previous, current, and next lessons.
    """
    prev_lesson, curr_lesson, next_lesson = None, None, None
    for i, line in enumerate(syllabus):
        if f"Lesson {current_lesson - 1}:" in line:
            prev_lesson = i
        elif f"Lesson {current_lesson}:" in line:
            curr_lesson = i
        elif f"Lesson {current_lesson + 1}:" in line:
            next_lesson = i
        elif f"Lesson {current_lesson + 2}:" in line:
            end_lesson = i
            break  # Stop after finding the next lesson to avoid unnecessary processing

    return prev_lesson, curr_lesson, next_lesson, end_lesson


# load objectives from syllabus
def extract_lesson_objectives(syllabus_path: Path, current_lesson: int, only_current: bool = False) -> str:
    """
    Extracts objectives for the previous, current, and next lessons from the syllabus.

    Args:
        syllabus_path (Path): The path to the Word document containing the syllabus.
        current_lesson (int): The current lesson number.

    Returns:
        str: The objectives for the previous, current, and next lessons.
    """
    max_retries = 3
    retry_delay = 10  # seconds

    for attempt in range(max_retries):
        try:
            doc = Document(str(syllabus_path))
            break
        except PackageNotFoundError:
            if attempt < max_retries - 1:
                print(f"Document `{syllabus_path.name}` is currently open. Retrying in {retry_delay} seconds...")
                time.sleep(retry_delay)
            else:
                raise PackageNotFoundError("Unable to open the document after multiple attempts. Please close the file and try again.")

    doc = Document(str(syllabus_path))
    syllabus_content = [para.text for para in doc.paragraphs]

    # Find the indices of the lessons
    prev_idx, curr_idx, next_idx, end_idx = find_lesson_indices(syllabus_content, current_lesson)

    # Extract the content for the previous, current, and next lessons
    prev_lesson_content = syllabus_content[prev_idx:curr_idx] if prev_idx is not None else []
    curr_lesson_content = syllabus_content[curr_idx:next_idx] if curr_idx is not None else []
    next_lesson_content = syllabus_content[next_idx:end_idx] if next_idx is not None else []

    # Combine the content
    combined_content = "\n".join(prev_lesson_content + curr_lesson_content + next_lesson_content)

    if only_current:
        return curr_lesson_content

    return combined_content


# # Load readings from a PDF file
# def load_readings(pdf_path: Path) -> str:
#     """
#     Loads text from a PDF file and returns it as a string.
#     The text is prefixed with the title derived from the PDF file name.

#     Args:
#         pdf_path (Path): The path to the PDF file.
#     Returns:
#         str: The text extracted from the PDF file.
#     Raises:
#         ValueError: If no readable text could be extracted from the PDF.
#     """
#     with open(str(pdf_path), 'rb') as file:
#         reader = PyPDF2.PdfReader(file)
#         text = 'title: ' + pdf_path.stem + "\n"
#         extracted_text = ""

#         for page in reader.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 extracted_text += page_text

#     # Check if any text was actually extracted
#     if not extracted_text.strip():  # Checks if extracted_text is empty or just whitespace
#         raise ValueError(f"No readable text found in {pdf_path.name}. Ensure the PDF is in a readable format.")

#     text += extracted_text
#     return text

# Load readings from either a PDF or a TXT file
def load_readings(file_path: Path) -> str:
    """
    Loads text from a PDF or TXT file and returns it as a string.
    The text is prefixed with the title derived from the file name.

    Args:
        file_path (Path): The path to the file.

    Returns:
        str: The text extracted from the file.

    Raises:
        ValueError: If no readable text could be extracted from the file.
    """

    # If the file is a PDF
    if file_path.suffix == '.pdf':
        with open(str(file_path), 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = 'title: ' + file_path.stem + "\n"
            extracted_text = ""

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    extracted_text += page_text

        # Check if any text was actually extracted
        if not extracted_text.strip():  # Checks if extracted_text is empty or just whitespace
            raise ValueError(f"No readable text found in {file_path.name}. Ensure the PDF is in a readable format.")

        text += extracted_text

    # If the file is a TXT file
    elif file_path.suffix == '.txt':
        text = 'title: ' + file_path.stem + "\n"  # Add the title from the filename
        extracted_text = ""

        # Try reading the file with utf-8 encoding, then fall back to ISO-8859-1
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                extracted_text = file.read()
        except UnicodeDecodeError:
            print(f"UnicodeDecodeError in {file_path.name}, using ISO-8859-1 instead")
            with open(file_path, 'r', encoding='ISO-8859-1') as file:
                extracted_text = file.read()
            print("Successfully loaded text using alternate method.")

        # Check if any text was actually extracted
        if not extracted_text.strip():
            raise ValueError(f"No readable text found in {file_path.name}. Ensure the TXT file has content.")

        text += extracted_text

    return text


def load_beamer_presentation(tex_path: Path) -> str:
    """
    Loads a Beamer presentation from a .tex file and returns it as a string.

    Args:
        tex_path (Path): The path to the .tex file containing the Beamer presentation.
    Returns:
        str: The content of the .tex file.
    """
    with open(tex_path, 'r', encoding='utf-8') as file:
        beamer_text = file.read()
    return beamer_text


# clean response
def clean_latex_content(latex_content: str) -> str:
    """
    Cleans LaTeX content by removing any text before the \title command and
    stripping extraneous LaTeX code blocks markers.

    Args:
        latex_content (str): The LaTeX content to be cleaned.

    Returns:
        str: The cleaned LaTeX content.
    """
    # Find the position of the \title command
    title_position = latex_content.find(r'\title')

    if title_position != -1:
        # Keep only the content starting from \title
        cleaned_content = latex_content[title_position:]
    else:
        # If \title is not found, return the original content (or handle as needed)
        cleaned_content = latex_content

    cleaned_content = cleaned_content.lstrip("```latex\n").rstrip("```")
    return cleaned_content
