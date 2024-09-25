import re
import time
# initial prompt
from pathlib import Path
from typing import Dict, List, Tuple, Union

import PyPDF2
from docx import Document
# doc import
from docx.opc.exceptions import PackageNotFoundError

from src.concept_web.tools import logger_setup


def load_documents(directory: Path, lesson_number) -> List[str]:
    """
    Load all documents from a single directory.

    Args:
        directory (Path): The directory to search for documents.

    Returns:
        List[str]: A list of document contents as strings.
    """
    all_documents = []

    # Iterate through files in the given directory (no recursion)
    for file in directory.glob('*'):
        if file.suffix in ['.pdf', '.txt', '.docx']:
            inferred_lesson_number = infer_lesson_from_filename(file.name)
            if inferred_lesson_number == lesson_number:
                all_documents.append(load_readings(file))

    return all_documents


def load_lessons(directories: Union[Path, List[Path]], lesson_range: range = None, recursive: bool = False, infer_from: str = "filename") -> List[str]:
    """
    Load specific lessons from one or multiple directories, with options to infer lesson numbers from filenames or directory names.

    Args:
        directories (Union[Path, List[Path]]): The directory or list of directories to search for documents.
        lesson_range (range, optional): The range of lesson numbers to load. If None, all lessons will be loaded.
        recursive (bool): If True, search through all subdirectories recursively.
        infer_from (str): Method to infer lesson numbers; options are "filename" or "directory".

    Returns:
        List[str]: A list of document contents as strings.
    """
    logger = logger_setup(log_level='WARNING')
    if isinstance(directories, (str, Path)):
        directories = [Path(directories)]

    all_documents = []

    for directory in directories:
        if recursive:
            # Use rglob for recursive search in all subdirectories
            subdirectories = [p for p in directory.rglob('*') if p.is_dir()]
            for subdirectory in subdirectories:
                inferred_lesson_number = infer_lesson_number(subdirectory, infer_from)
                if lesson_range is None or inferred_lesson_number in lesson_range:
                    all_documents.extend(load_documents(subdirectory, inferred_lesson_number))
                    # Check if any subdirectory has its own subdirectory
                else:
                    continue
                sub_subdirectories = [p for p in subdirectory.glob('*/') if p.is_dir()]
                if sub_subdirectories:
                    logger.warning(f"Overly nested lesson directories found: {subdirectory} contains subdirectories. Readings in these directories not loaded")

        else:
            inferred_lesson_number = infer_lesson_number(directory, infer_from)
            if lesson_range is None or inferred_lesson_number in lesson_range:
                all_documents.extend(load_documents(directory, inferred_lesson_number))

    return all_documents


def infer_lesson_number(path: Path, infer_from: str) -> int:
    """
    Infers the lesson number from either the directory name or the filename.

    Args:
        path (Path): The directory or file path.
        infer_from (str): Method to infer lesson numbers; options are "filename" or "directory".

    Returns:
        int: The inferred lesson number.
    """
    if infer_from == "filename":
        return infer_lesson_from_filename(path.name)
    elif infer_from == "directory":
        # Assuming directory name contains the lesson number
        match = re.search(r'\d+', path.name)
        if match:
            return int(match.group(0))
    return None


def infer_lesson_from_filename(filename: str) -> int:
    # Adjust the regex to correctly match different patterns
    match = re.search(r'(?:Lesson|L)?(\d+)\.\d+|Lesson\s*(\d+)|L\s*(\d+)', filename, re.IGNORECASE)

    if match:
        # Return the first group that is not None
        for group in match.groups():
            if group:
                return int(group)
    return None


def extract_text_from_pdf(pdf_path: Union[str, Path]) -> str:
    """
    Extracts text from a PDF file, treating paragraph breaks correctly.

    Args:
        pdf_path (Path): The path to the PDF file.

    Returns:
        str: The text content of the PDF as a single string.
    """
    pdf_path = Path(pdf_path)

    text_content = []
    with open(str(pdf_path), 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                # Split the text into paragraphs (double newlines as paragraph breaks)
                paragraphs = page_text.split('\n\n')
                text_content.extend(paragraph.strip() for paragraph in paragraphs if paragraph.strip())
    return ' '.join(text_content)


def extract_lessons_from_page(page_content: str, lesson_marker: str = "Week") -> Dict[int, str]:
    """
    Extracts lessons and their content from a single page of syllabus content.

    Args:
        page_content (str): The content of the syllabus page as a string.
        lesson_marker (str): The text that marks the beginning of a lesson (default is "Week").

    Returns:
        Dict[int, str]: A dictionary where the keys are lesson numbers and the values are the corresponding content.
    """
    lessons = {}

    lesson_marker_not_used = 'Lesson' if lesson_marker == 'Week' else 'Week'
    lesson_pattern = re.compile(rf"({lesson_marker}\s*(\d+).*?)(?={lesson_marker}\s*\d+|$)", re.DOTALL)

    matches = lesson_pattern.finditer(page_content)

    if not matches:
        print(f"Trying alternative lesson indicator: {lesson_marker_not_used}")
        lesson_pattern = re.compile(rf"({lesson_marker_not_used}\s*(\d+).*?)(?={lesson_marker_not_used}\s*\d+|$)", re.DOTALL)
        matches = lesson_pattern.finditer(page_content)
        if not matches:
            raise ValueError("Unable to extract lesson objectives from syllabus. Consider manually entering a list of lesson objectives")

    for match in matches:
        full_match = match.group(1).strip()
        lesson_number = int(match.group(2))

        if ':' in full_match:
            content = full_match.split(":", 1)[1].strip()
        else:
            content = full_match[len(lesson_marker) + len(str(lesson_number)):].strip()

        lessons[f"Lesson {lesson_number}"] = content

    return lessons


def find_pdf_lessons(syllabus: dict, current_lesson: int) -> Tuple[str, str, str, str]:
    """
    Finds the lessons in the syllabus content based on the current lesson number for PDFs.

    Args:
        syllabus (dict): Dictionary of lesson content extracted from PDF.
        current_lesson (int): The lesson number for which to find surrounding lessons.

    Returns:
        Tuple[str, str, str, str]: The content for the previous, current, next, and the end of the next lesson.
    """
    prev_lesson = syllabus.get(f"Lesson {current_lesson-1}", "")
    curr_lesson = syllabus.get(f"Lesson {current_lesson}", "")
    next_lesson = syllabus.get(f"Lesson {current_lesson+1}", "")
    end_lesson = syllabus.get(f"Lesson {current_lesson+2}", "")

    return prev_lesson, curr_lesson, next_lesson, end_lesson


def find_docx_indices(syllabus: List[str], current_lesson: int, lesson_identifier: str = "Lesson") -> Tuple[int, int, int, int]:
    """
    Finds the indices of the lessons in the syllabus content.

    Args:
        syllabus (List[str]): A list of strings where each string represents a line in the syllabus document.
        current_lesson (int): The lesson number for which to find surrounding lessons.

    Returns:
        Tuple[int, int, int, int]: The indices of the previous, current, next, and the end of the next lesson.
    """
    prev_lesson, curr_lesson, next_lesson, end_lesson = None, None, None, None
    lesson_pattern = re.compile(rf"{lesson_identifier}\s*{current_lesson}.*?:")

    for i, line in enumerate(syllabus):
        if re.search(rf"{lesson_identifier}\s*{current_lesson - 1}.*?:?", line):
            prev_lesson = i
        elif lesson_pattern.search(line):
            curr_lesson = i
        elif re.search(rf"{lesson_identifier}\s*{current_lesson + 1}.*?:?", line):
            next_lesson = i
        elif re.search(rf"{lesson_identifier}\s*{current_lesson + 2}.*?:?", line):
            end_lesson = i
            break

    return prev_lesson, curr_lesson, next_lesson, end_lesson


def load_docx_syllabus(syllabus_path: Union[str, Path]) -> List[str]:
    """
    Loads a DOCX syllabus and returns its content as a list of paragraphs.

    Args:
        syllabus_path (Path): The path to the DOCX file.

    Returns:
        List[str]: The syllabus content as a list of paragraphs.
    """
    syllabus_path = Path(syllabus_path)

    max_retries = 3
    retry_delay = 10  # seconds

    for attempt in range(max_retries):
        try:
            doc = Document(str(syllabus_path))
            return [para.text for para in doc.paragraphs]
        except PackageNotFoundError:
            if attempt < max_retries - 1:
                print(f"Document `{syllabus_path.name}` is currently open. Retrying in {retry_delay} seconds...")
                time.sleep(retry_delay)
            else:
                raise PackageNotFoundError("Unable to open the document after multiple attempts. Please close the file and try again.")


def extract_lesson_objectives(syllabus_path: Union[str, Path], current_lesson: int, only_current: bool = False) -> str:
    """
    Extracts objectives for the previous, current, and next lessons from the syllabus.

    Args:
        syllabus_path (Path): The path to the syllabus document (PDF or DOCX).
        current_lesson (int): The current lesson number.

    Returns:
        str: The objectives for the previous, current, and next lessons.
    """
    syllabus_path = Path(syllabus_path)

    if syllabus_path.suffix == '.docx':
        syllabus_content = load_docx_syllabus(syllabus_path)
        prev_idx, curr_idx, next_idx, end_idx = find_docx_indices(syllabus_content, current_lesson)
        prev_lesson_content = "\n".join(syllabus_content[prev_idx:curr_idx]) if prev_idx is not None else ""
        curr_lesson_content = "\n".join(syllabus_content[curr_idx:next_idx]) if curr_idx is not None else ""
        next_lesson_content = "\n".join(syllabus_content[next_idx:end_idx]) if next_idx is not None else ""
    elif syllabus_path.suffix == '.pdf':
        page_content = extract_text_from_pdf(syllabus_path)
        lessons = extract_lessons_from_page(page_content)
        prev_lesson_content, curr_lesson_content, next_lesson_content, _ = find_pdf_lessons(lessons, current_lesson)
    else:
        raise ValueError(f"Unsupported file type: {syllabus_path.suffix}")

    combined_content = "\n".join(filter(None, [prev_lesson_content, curr_lesson_content, next_lesson_content]))

    if only_current:
        return curr_lesson_content

    return combined_content


# Load readings from either a PDF or a TXT file
def load_readings(file_path: Union[str, Path]) -> str:
    """
    Loads text from a PDF, DOCX, or TXT file and returns it as a string.
    The text is prefixed with the title derived from the file name.

    Args:
        file_path (Path): The path to the file.

    Returns:
        str: The text extracted from the file.

    Raises:
        ValueError: If no readable text could be extracted from the file.
    """
    file_path = Path(file_path
                     )

    def check_extracted_text(extracted_text: str, file_name: str):
        if not extracted_text.strip():
            raise ValueError(f"No readable text found in {file_name}. Ensure the file has content.")

    text = 'title: ' + file_path.stem + "\n"

    if file_path.suffix.lower() == '.pdf':
        extracted_text = extract_text_from_pdf(file_path)
        check_extracted_text(extracted_text, file_path.name)

    elif file_path.suffix.lower() == '.docx':
        try:
            doc = Document(str(file_path))
            extracted_text = "\n".join([para.text for para in doc.paragraphs])
        except PackageNotFoundError:
            raise ValueError(f"Unable to open {file_path.name}. The file might be corrupted or not a valid DOCX document.")
        check_extracted_text(extracted_text, file_path.name)

    elif file_path.suffix.lower() == '.txt':
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                extracted_text = file.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='ISO-8859-1') as file:
                extracted_text = file.read()
        check_extracted_text(extracted_text, file_path.name)

    else:
        raise ValueError(f"Unsupported file type: {file_path.suffix}")

    return text + extracted_text


if __name__ == "__main__":
    import os

    from dotenv import load_dotenv
    # llm chain setup
    from langchain_openai import ChatOpenAI
    # self-defined utils
    from load_documents import extract_lesson_objectives, load_readings
    load_dotenv()

    # Path definitions
    syllabus_path = Path(os.getenv('syllabus_path'))
    pdf_syllabus_path = Path(os.getenv('pdf_syllabus_path'))
    readingsDir = Path(os.getenv('readingsDir'))

    lsn = 8
    lsn_objectives_doc = extract_lesson_objectives(syllabus_path,
                                                   lsn,
                                                   only_current=True)

    lsn_objectives_pdf = extract_lesson_objectives(pdf_syllabus_path,
                                                   lsn,
                                                   only_current=True)

    print(f"doc objectives:\n{lsn_objectives_doc}")
    print(f"\npdf objectives:\n{lsn_objectives_pdf}")

    docs = load_lessons(readingsDir, recursive=True, lesson_range=range(1, 4))
